#!/usr/bin/env python3
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_code_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    font = run.font
    font.name = 'Courier New'
    # Ensure proper font on macOS/Windows
    try:
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Courier New')
        rFonts.set(qn('w:hAnsi'), 'Courier New')
        rPr.append(rFonts)
    except Exception:
        pass


def convert(md_path: Path, out_path: Path):
    text = md_path.read_text(encoding='utf-8')
    doc = Document()

    in_code = False
    code_lines = []

    for raw_line in text.splitlines():
        line = raw_line.rstrip('\n')
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_code_paragraph(doc, '\n'.join(code_lines))
                code_lines = []
            continue

        if in_code:
            code_lines.append(raw_line)
            continue

        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
            continue

        # Lists
        if line.lstrip().startswith('- '):
            doc.add_paragraph(line.lstrip()[2:].strip(), style='List Bullet')
            continue

        # Blank lines
        if line.strip() == '':
            doc.add_paragraph('')
            continue

        # Regular paragraph (very simple inline handling: strip backticks)
        paragraph_text = line.replace('`', '')
        doc.add_paragraph(paragraph_text)

    # If file ended inside a code block
    if in_code and code_lines:
        add_code_paragraph(doc, '\n'.join(code_lines))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print('Usage: md_to_docx.py <input.md> <output.docx>')
        sys.exit(2)
    inp = Path(sys.argv[1]).expanduser().resolve()
    outp = Path(sys.argv[2]).expanduser()
    if not inp.exists():
        print(f'Input file not found: {inp}')
        sys.exit(1)
    convert(inp, outp)
    print(f'Wrote {outp}')
